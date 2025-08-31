from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Any, Optional, List
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import re
from datetime import datetime

class DocumentMetadata:
    def __init__(self):
        self.title: Optional[str] = None
        self.author: Optional[str] = None
        self.creation_date: Optional[datetime] = None
        self.modification_date: Optional[datetime] = None
        self.page_count: Optional[int] = None
        self.word_count: Optional[int] = None
        self.character_count: Optional[int] = None

class DocumentProcessor(ABC):
    @abstractmethod
    def extract_text(self, file_path: Path) -> str:
        pass
    
    @abstractmethod
    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        pass

class PDFProcessor(DocumentProcessor):
    def extract_text(self, file_path: Path) -> str:
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text.strip()
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        metadata = DocumentMetadata()
        try:
            doc = fitz.open(file_path)
            pdf_metadata = doc.metadata
            
            metadata.title = pdf_metadata.get('title')
            metadata.author = pdf_metadata.get('author')
            metadata.page_count = doc.page_count
            
            # Extract text to count words
            text = self.extract_text(file_path)
            metadata.word_count = len(text.split())
            metadata.character_count = len(text)
            
            doc.close()
        except Exception:
            pass  # Return partial metadata
        
        return metadata

class DOCXProcessor(DocumentProcessor):
    def extract_text(self, file_path: Path) -> str:
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        metadata = DocumentMetadata()
        try:
            doc = DocxDocument(file_path)
            core_props = doc.core_properties
            
            metadata.title = core_props.title
            metadata.author = core_props.author
            metadata.creation_date = core_props.created
            metadata.modification_date = core_props.modified
            
            # Extract text to count words and characters
            text = self.extract_text(file_path)
            metadata.word_count = len(text.split())
            metadata.character_count = len(text)
            
        except Exception:
            pass  # Return partial metadata
        
        return metadata

class TextProcessor(DocumentProcessor):
    def extract_text(self, file_path: Path) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read().strip()
        except Exception as e:
            raise Exception(f"Error processing text file: {str(e)}")
    
    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        metadata = DocumentMetadata()
        try:
            text = self.extract_text(file_path)
            metadata.word_count = len(text.split())
            metadata.character_count = len(text)
            
            # Get file stats
            stat = file_path.stat()
            metadata.creation_date = datetime.fromtimestamp(stat.st_ctime)
            metadata.modification_date = datetime.fromtimestamp(stat.st_mtime)
            
        except Exception:
            pass
        
        return metadata

class MarkdownProcessor(TextProcessor):
    def extract_text(self, file_path: Path) -> str:
        # For markdown, we could strip markdown syntax but for now keep it
        return super().extract_text(file_path)

class DocumentProcessorFactory:
    @staticmethod
    def get_processor(mime_type: str) -> DocumentProcessor:
        processors = {
            "application/pdf": PDFProcessor(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DOCXProcessor(),
            "text/plain": TextProcessor(),
            "text/markdown": MarkdownProcessor()
        }
        
        processor = processors.get(mime_type)
        if not processor:
            raise ValueError(f"Unsupported file type: {mime_type}")
        
        return processor

def validate_file_type(filename: str, mime_type: str) -> bool:
    """Validate file type based on filename and MIME type"""
    allowed_extensions = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
        ".md": "text/markdown"
    }
    
    file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
    file_ext = f'.{file_ext}'
    
    return file_ext in allowed_extensions and allowed_extensions[file_ext] == mime_type

def validate_file_size(file_size: int, max_size: int = 100 * 1024 * 1024) -> bool:
    """Validate file size (default max: 100MB)"""
    return 0 < file_size <= max_size

def chunk_text(text: str, chunk_size: int = 1000, overlap_percentage: int = 10) -> List[Dict[str, Any]]:
    """Split text into overlapping chunks"""
    if not text:
        return []
    
    overlap_size = int(chunk_size * overlap_percentage / 100)
    chunks = []
    start = 0
    chunk_id = 0
    
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk_text = text[start:end]
        
        chunks.append({
            "chunk_id": f"chunk_{chunk_id}",
            "text": chunk_text,
            "start_char": start,
            "end_char": end,
            "size": len(chunk_text)
        })
        
        # Move start position with overlap
        start = end - overlap_size
        if start >= end:  # Prevent infinite loop
            break
            
        chunk_id += 1
    
    return chunks