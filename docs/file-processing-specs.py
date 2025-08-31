# DeepInsight File Processing & Document Parsing Specifications
# Complete document processing pipeline for PDF, DOCX, TXT, and MD files

from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import magic
import mimetypes
from dataclasses import dataclass
import hashlib
import os

# Document processing libraries
try:
    import pymupdf as fitz  # PyMuPDF for PDF processing
except ImportError:
    import fitz

import docx  # python-docx for Word documents
import markdown  # For markdown files
import re
from datetime import datetime

# ============================================================================
# Configuration and Constants
# ============================================================================

# Allowed MIME types and file extensions
ALLOWED_MIME_TYPES = {
    'pdf': 'application/pdf',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'txt': 'text/plain',
    'md': 'text/markdown'
}

ALLOWED_EXTENSIONS = {
    '.pdf': 'application/pdf',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    '.txt': 'text/plain',
    '.md': 'text/markdown'
}

# File size limits
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
MIN_FILE_SIZE = 1  # 1 byte

# Text processing limits
MAX_TEXT_LENGTH = 10 * 1024 * 1024  # 10MB of text
MIN_TEXT_LENGTH = 10  # 10 characters minimum

# ============================================================================
# Data Classes
# ============================================================================

@dataclass
class DocumentMetadata:
    """Metadata extracted from document"""
    filename: str
    file_size: int
    mime_type: str
    file_hash: str
    created_at: Optional[datetime] = None
    modified_at: Optional[datetime] = None
    author: Optional[str] = None
    title: Optional[str] = None
    subject: Optional[str] = None
    creator: Optional[str] = None
    producer: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None
    language: Optional[str] = None

@dataclass
class ProcessedDocument:
    """Complete processed document with text and metadata"""
    text: str
    metadata: DocumentMetadata
    pages: Optional[List[str]] = None  # Text per page (for PDFs)
    structure: Optional[Dict[str, Any]] = None  # Document structure info

@dataclass
class ValidationResult:
    """File validation result"""
    is_valid: bool
    error_message: Optional[str] = None
    warnings: List[str] = None

# ============================================================================
# File Validation Functions
# ============================================================================

class FileValidator:
    """Comprehensive file validation for security and format checking"""
    
    @staticmethod
    def validate_file_size(file_path: Path) -> ValidationResult:
        """Validate file size is within allowed limits"""
        try:
            file_size = file_path.stat().st_size
            
            if file_size < MIN_FILE_SIZE:
                return ValidationResult(
                    is_valid=False,
                    error_message="File is empty or corrupted"
                )
            
            if file_size > MAX_FILE_SIZE:
                return ValidationResult(
                    is_valid=False,
                    error_message=f"File size ({file_size:,} bytes) exceeds maximum allowed size ({MAX_FILE_SIZE:,} bytes)"
                )
            
            return ValidationResult(is_valid=True)
            
        except Exception as e:
            return ValidationResult(
                is_valid=False,
                error_message=f"Could not read file size: {str(e)}"
            )
    
    @staticmethod
    def validate_file_type(file_path: Path) -> ValidationResult:
        """Validate file type using multiple methods for security"""
        filename = file_path.name
        file_extension = file_path.suffix.lower()
        warnings = []
        
        # Check file extension
        if file_extension not in ALLOWED_EXTENSIONS:
            return ValidationResult(
                is_valid=False,
                error_message=f"File extension '{file_extension}' is not allowed. Allowed extensions: {list(ALLOWED_EXTENSIONS.keys())}"
            )
        
        expected_mime_type = ALLOWED_EXTENSIONS[file_extension]
        
        try:
            # Use python-magic for MIME type detection
            detected_mime_type = magic.from_file(str(file_path), mime=True)
            
            # Check if detected MIME type matches expected
            if detected_mime_type != expected_mime_type:
                # Some acceptable variations
                acceptable_variations = {
                    'text/plain': ['text/plain', 'text/x-python', 'text/x-script'],
                    'text/markdown': ['text/markdown', 'text/x-markdown', 'text/plain'],
                    'application/pdf': ['application/pdf'],
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': [
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        'application/zip'  # DOCX files are ZIP archives
                    ]
                }
                
                allowed_types = acceptable_variations.get(expected_mime_type, [expected_mime_type])
                
                if detected_mime_type not in allowed_types:
                    return ValidationResult(
                        is_valid=False,
                        error_message=f"File content type '{detected_mime_type}' does not match extension '{file_extension}'. Expected: {expected_mime_type}"
                    )
                else:
                    warnings.append(f"MIME type variation detected: {detected_mime_type} (expected: {expected_mime_type})")
            
            return ValidationResult(
                is_valid=True,
                warnings=warnings
            )
            
        except Exception as e:
            # Fallback to mimetypes module
            try:
                guessed_type, _ = mimetypes.guess_type(str(file_path))
                if guessed_type != expected_mime_type:
                    warnings.append(f"Could not verify MIME type with magic, using mimetypes: {guessed_type}")
                
                return ValidationResult(
                    is_valid=True,
                    warnings=warnings
                )
            except Exception as fallback_error:
                return ValidationResult(
                    is_valid=False,
                    error_message=f"Could not determine file type: {str(e)}"
                )
    
    @staticmethod
    def validate_file_content(file_path: Path) -> ValidationResult:
        """Validate file content for basic integrity"""
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return FileValidator._validate_pdf_content(file_path)
            elif file_extension == '.docx':
                return FileValidator._validate_docx_content(file_path)
            elif file_extension in ['.txt', '.md']:
                return FileValidator._validate_text_content(file_path)
            else:
                return ValidationResult(
                    is_valid=False,
                    error_message=f"Unsupported file type: {file_extension}"
                )
                
        except Exception as e:
            return ValidationResult(
                is_valid=False,
                error_message=f"Content validation failed: {str(e)}"
            )
    
    @staticmethod
    def _validate_pdf_content(file_path: Path) -> ValidationResult:
        """Validate PDF file content"""
        try:
            doc = fitz.open(str(file_path))
            
            if doc.page_count == 0:
                return ValidationResult(
                    is_valid=False,
                    error_message="PDF file contains no pages"
                )
            
            # Check if PDF is password protected
            if doc.needs_pass:
                return ValidationResult(
                    is_valid=False,
                    error_message="PDF file is password protected and cannot be processed"
                )
            
            # Try to extract text from first page to verify readability
            first_page = doc[0]
            text = first_page.get_text()
            
            doc.close()
            
            return ValidationResult(is_valid=True)
            
        except Exception as e:
            return ValidationResult(
                is_valid=False,
                error_message=f"PDF validation failed: {str(e)}"
            )
    
    @staticmethod
    def _validate_docx_content(file_path: Path) -> ValidationResult:
        """Validate DOCX file content"""
        try:
            doc = docx.Document(str(file_path))
            
            # Try to access document properties
            _ = doc.core_properties.title
            
            # Try to access paragraphs
            paragraphs = doc.paragraphs
            
            return ValidationResult(is_valid=True)
            
        except Exception as e:
            return ValidationResult(
                is_valid=False,
                error_message=f"DOCX validation failed: {str(e)}"
            )
    
    @staticmethod
    def _validate_text_content(file_path: Path) -> ValidationResult:
        """Validate text file content"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                # Try to read the first 1KB to check encoding
                content = f.read(1024)
                
                if len(content) == 0:
                    return ValidationResult(
                        is_valid=False,
                        error_message="Text file is empty"
                    )
            
            return ValidationResult(is_valid=True)
            
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        f.read(1024)
                    return ValidationResult(
                        is_valid=True,
                        warnings=[f"File encoding detected as {encoding} instead of UTF-8"]
                    )
                except UnicodeDecodeError:
                    continue
            
            return ValidationResult(
                is_valid=False,
                error_message="Text file has unreadable encoding"
            )
        except Exception as e:
            return ValidationResult(
                is_valid=False,
                error_message=f"Text file validation failed: {str(e)}"
            )
    
    @classmethod
    def validate_file(cls, file_path: Path) -> ValidationResult:
        """Comprehensive file validation"""
        # Check if file exists
        if not file_path.exists():
            return ValidationResult(
                is_valid=False,
                error_message="File does not exist"
            )
        
        if not file_path.is_file():
            return ValidationResult(
                is_valid=False,
                error_message="Path is not a file"
            )
        
        # Validate file size
        size_result = cls.validate_file_size(file_path)
        if not size_result.is_valid:
            return size_result
        
        # Validate file type
        type_result = cls.validate_file_type(file_path)
        if not type_result.is_valid:
            return type_result
        
        # Validate content
        content_result = cls.validate_file_content(file_path)
        if not content_result.is_valid:
            return content_result
        
        # Combine warnings
        all_warnings = []
        for result in [size_result, type_result, content_result]:
            if result.warnings:
                all_warnings.extend(result.warnings)
        
        return ValidationResult(
            is_valid=True,
            warnings=all_warnings if all_warnings else None
        )

# ============================================================================
# Document Processor Base Class
# ============================================================================

class DocumentProcessor(ABC):
    """Abstract base class for document processors"""
    
    @abstractmethod
    def extract_text(self, file_path: Path) -> str:
        """Extract text from document"""
        pass
    
    @abstractmethod
    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract document metadata"""
        pass
    
    def process_document(self, file_path: Path) -> ProcessedDocument:
        """Process document to extract text and metadata"""
        text = self.extract_text(file_path)
        metadata = self.extract_metadata(file_path)
        
        # Add word count and character count to metadata
        metadata.word_count = len(text.split())
        metadata.character_count = len(text)
        
        return ProcessedDocument(
            text=text,
            metadata=metadata
        )
    
    @staticmethod
    def calculate_file_hash(file_path: Path) -> str:
        """Calculate SHA-256 hash of file"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()

# ============================================================================
# Specific Document Processors
# ============================================================================

class PDFProcessor(DocumentProcessor):
    """PDF document processor using PyMuPDF"""
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from PDF using PyMuPDF"""
        doc = fitz.open(str(file_path))
        text_content = []
        
        try:
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                
                # Clean up the text
                page_text = self._clean_pdf_text(page_text)
                text_content.append(page_text)
            
            return "\n\n".join(text_content)
            
        finally:
            doc.close()
    
    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract PDF metadata"""
        doc = fitz.open(str(file_path))
        
        try:
            metadata_dict = doc.metadata
            file_stats = file_path.stat()
            
            # Convert PyMuPDF dates (if present)
            created_at = None
            modified_at = None
            
            if metadata_dict.get('creationDate'):
                try:
                    created_at = self._parse_pdf_date(metadata_dict['creationDate'])
                except:
                    pass
            
            if metadata_dict.get('modDate'):
                try:
                    modified_at = self._parse_pdf_date(metadata_dict['modDate'])
                except:
                    pass
            
            return DocumentMetadata(
                filename=file_path.name,
                file_size=file_stats.st_size,
                mime_type='application/pdf',
                file_hash=self.calculate_file_hash(file_path),
                created_at=created_at,
                modified_at=modified_at or datetime.fromtimestamp(file_stats.st_mtime),
                author=metadata_dict.get('author'),
                title=metadata_dict.get('title'),
                subject=metadata_dict.get('subject'),
                creator=metadata_dict.get('creator'),
                producer=metadata_dict.get('producer'),
                page_count=doc.page_count
            )
            
        finally:
            doc.close()
    
    def extract_pages(self, file_path: Path) -> List[str]:
        """Extract text per page"""
        doc = fitz.open(str(file_path))
        pages = []
        
        try:
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                cleaned_text = self._clean_pdf_text(page_text)
                pages.append(cleaned_text)
            
            return pages
            
        finally:
            doc.close()
    
    @staticmethod
    def _clean_pdf_text(text: str) -> str:
        """Clean extracted PDF text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page numbers and headers/footers (simple heuristic)
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            # Skip very short lines that might be page numbers
            if len(line) > 3:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    @staticmethod
    def _parse_pdf_date(date_str: str) -> datetime:
        """Parse PDF date string to datetime"""
        # PDF dates are in format: D:YYYYMMDDHHmmSSOHH'mm
        if date_str.startswith('D:'):
            date_str = date_str[2:]
        
        # Extract just the date part (first 14 characters: YYYYMMDDHHMMSS)
        date_part = date_str[:14]
        return datetime.strptime(date_part, '%Y%m%d%H%M%S')

class DOCXProcessor(DocumentProcessor):
    """DOCX document processor using python-docx"""
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from DOCX using python-docx"""
        doc = docx.Document(str(file_path))
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            table_text = self._extract_table_text(table)
            if table_text:
                text_content.append(table_text)
        
        return '\n\n'.join(text_content)
    
    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract DOCX metadata"""
        doc = docx.Document(str(file_path))
        props = doc.core_properties
        file_stats = file_path.stat()
        
        return DocumentMetadata(
            filename=file_path.name,
            file_size=file_stats.st_size,
            mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            file_hash=self.calculate_file_hash(file_path),
            created_at=props.created,
            modified_at=props.modified or datetime.fromtimestamp(file_stats.st_mtime),
            author=props.author,
            title=props.title,
            subject=props.subject,
            language=props.language
        )
    
    @staticmethod
    def _extract_table_text(table) -> str:
        """Extract text from a DOCX table"""
        table_text = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            
            if row_text:
                table_text.append(' | '.join(row_text))
        
        return '\n'.join(table_text)

class TextProcessor(DocumentProcessor):
    """Plain text processor"""
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    # Clean up the text
                    return self._clean_text(content)
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode text file with any supported encoding")
    
    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract text file metadata"""
        file_stats = file_path.stat()
        
        return DocumentMetadata(
            filename=file_path.name,
            file_size=file_stats.st_size,
            mime_type='text/plain',
            file_hash=self.calculate_file_hash(file_path),
            modified_at=datetime.fromtimestamp(file_stats.st_mtime)
        )
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean extracted text"""
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive whitespace but preserve paragraph structure
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            cleaned_line = ' '.join(line.split())  # Remove extra spaces
            cleaned_lines.append(cleaned_line)
        
        # Preserve paragraph breaks but remove excessive empty lines
        result = '\n'.join(cleaned_lines)
        result = re.sub(r'\n{3,}', '\n\n', result)  # Max 2 consecutive newlines
        
        return result.strip()

class MarkdownProcessor(DocumentProcessor):
    """Markdown document processor"""
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from markdown file"""
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    markdown_content = f.read()
                
                # Convert markdown to plain text
                md = markdown.Markdown()
                html = md.convert(markdown_content)
                
                # Remove HTML tags to get plain text
                plain_text = re.sub(r'<[^>]+>', '', html)
                
                # Clean up the text
                return self._clean_markdown_text(plain_text)
                
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode markdown file with any supported encoding")
    
    def extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract markdown file metadata"""
        file_stats = file_path.stat()
        
        # Try to extract title from first heading
        title = None
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                first_lines = f.read(1000)  # Read first 1KB
                
                # Look for first H1 heading
                h1_match = re.search(r'^#\s+(.+)$', first_lines, re.MULTILINE)
                if h1_match:
                    title = h1_match.group(1).strip()
        except:
            pass
        
        return DocumentMetadata(
            filename=file_path.name,
            file_size=file_stats.st_size,
            mime_type='text/markdown',
            file_hash=self.calculate_file_hash(file_path),
            modified_at=datetime.fromtimestamp(file_stats.st_mtime),
            title=title
        )
    
    @staticmethod
    def _clean_markdown_text(text: str) -> str:
        """Clean extracted markdown text"""
        # Decode HTML entities
        text = text.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
        
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Normalize line endings and remove excessive newlines
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()

# ============================================================================
# Document Processing Factory
# ============================================================================

class DocumentProcessorFactory:
    """Factory class for creating appropriate document processors"""
    
    _processors = {
        '.pdf': PDFProcessor,
        '.docx': DOCXProcessor,
        '.txt': TextProcessor,
        '.md': MarkdownProcessor
    }
    
    @classmethod
    def get_processor(cls, file_path: Path) -> DocumentProcessor:
        """Get appropriate processor for file type"""
        file_extension = file_path.suffix.lower()
        
        if file_extension not in cls._processors:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        processor_class = cls._processors[file_extension]
        return processor_class()
    
    @classmethod
    def process_file(cls, file_path: Path) -> Tuple[ProcessedDocument, ValidationResult]:
        """Complete file processing with validation"""
        # Validate file first
        validation_result = FileValidator.validate_file(file_path)
        
        if not validation_result.is_valid:
            return None, validation_result
        
        # Process file
        try:
            processor = cls.get_processor(file_path)
            processed_doc = processor.process_document(file_path)
            
            # Additional text validation
            if len(processed_doc.text) < MIN_TEXT_LENGTH:
                return None, ValidationResult(
                    is_valid=False,
                    error_message=f"Extracted text is too short ({len(processed_doc.text)} characters). Minimum required: {MIN_TEXT_LENGTH}"
                )
            
            if len(processed_doc.text) > MAX_TEXT_LENGTH:
                return None, ValidationResult(
                    is_valid=False,
                    error_message=f"Extracted text is too long ({len(processed_doc.text)} characters). Maximum allowed: {MAX_TEXT_LENGTH}"
                )
            
            return processed_doc, validation_result
            
        except Exception as e:
            return None, ValidationResult(
                is_valid=False,
                error_message=f"Document processing failed: {str(e)}"
            )

# ============================================================================
# Utility Functions
# ============================================================================

def get_supported_file_types() -> Dict[str, str]:
    """Get dictionary of supported file extensions and their MIME types"""
    return ALLOWED_EXTENSIONS.copy()

def is_supported_file_type(filename: str) -> bool:
    """Check if filename has supported extension"""
    extension = Path(filename).suffix.lower()
    return extension in ALLOWED_EXTENSIONS

def estimate_processing_time(file_size: int) -> float:
    """Estimate processing time in seconds based on file size"""
    # Rough estimates based on file size (in seconds)
    if file_size < 1024 * 1024:  # < 1MB
        return 5.0
    elif file_size < 10 * 1024 * 1024:  # < 10MB
        return 30.0
    elif file_size < 50 * 1024 * 1024:  # < 50MB
        return 120.0
    else:  # >= 50MB
        return 300.0

def sanitize_filename(filename: str) -> str:
    """Sanitize filename for safe storage"""
    # Remove or replace dangerous characters
    sanitized = re.sub(r'[<>:"/\\|?*]', '_', filename)
    
    # Remove leading/trailing spaces and dots
    sanitized = sanitized.strip(' .')
    
    # Limit length
    if len(sanitized) > 255:
        name, ext = os.path.splitext(sanitized)
        max_name_length = 255 - len(ext)
        sanitized = name[:max_name_length] + ext
    
    return sanitized